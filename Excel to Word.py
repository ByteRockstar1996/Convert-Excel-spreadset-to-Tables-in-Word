from docx import Document
import xlrd
import sys

def process_excel(excel_file, doc_file):

    document = Document()

    workbook = xlrd.open_workbook(excel_file)
    headings = None

    sheet = workbook.sheet_by_index(0)
    for row in sheet.get_rows():

        if headings == None:
            headings = [x.value.strip() for x in row]
        else :
            table = document.add_table(rows=len(headings), cols=2)
            table.columns[0].width=2000000
            table.columns[1].width=4000000
            table.allow_autofit=True
            table.style='Table Grid'
            
            for j, col in enumerate(headings):
                cell = table.cell(j, 0)
                cell.text = col
                cell = table.cell(j, 1)
                cell.text = row[j].value
            
            bolding_columns = [0]
            for row in list(range(len(headings))):
                for column in bolding_columns:
                    table.rows[row].cells[column].paragraphs[0].runs[0].font.bold = True

    document.save(doc_file)
    
ExcelFile=sys.argv[1]
process_excel(ExcelFile, ExcelFile[:-4]+'docx')
